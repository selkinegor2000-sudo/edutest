"""
Конвертирует docs/practical_tables_and_diagrams.md в Word-файл
с настоящими таблицами (не символами |), готовый для вставки в диплом.
"""
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
INPUT = os.path.join(BASE, "docs", "practical_tables_and_diagrams.md")
OUTPUT = os.path.join(BASE, "docs", "ТАБЛИЦЫ_ДЛЯ_ДИПЛОМА.docx")


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table_from_md(doc, lines):
    # lines — строки markdown-таблицы (убираем строку-разделитель ---)
    rows = [l for l in lines if not re.match(r"^\s*\|?\s*[-:]+[-| :]*\s*$", l)]
    if not rows:
        return
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    col_count = max(len(r) for r in parsed)
    tbl = doc.add_table(rows=len(parsed), cols=col_count)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(parsed):
        for j, cell_text in enumerate(row):
            if j >= col_count:
                break
            cell = tbl.rows[i].cells[j]
            cell.text = cell_text
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.runs[0] if para.runs else para.add_run(cell_text)
            run.font.size = Pt(10)
            if i == 0:
                run.font.bold = True
                set_cell_bg(cell, "D9E1F2")

    doc.add_paragraph()


def main():
    with open(INPUT, encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Стиль Normal
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    lines = content.split("\n")
    i = 0
    in_code = False

    while i < len(lines):
        line = lines[i]

        # Пропускаем блоки кода (mermaid и т.д.)
        if line.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            i += 1
            continue

        # Заголовки
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            title = line[3:].strip()
            p = doc.add_heading(title, level=2)
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Таблица markdown
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table_from_md(doc, table_lines)
            continue

        # Обычный абзац
        stripped = line.strip()
        if stripped:
            p = doc.add_paragraph(stripped)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        i += 1

    doc.save(OUTPUT)
    print(f"Сохранено: {OUTPUT}")


if __name__ == "__main__":
    main()
